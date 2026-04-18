"""
Generate a professional DOCX report for Vault-AI project.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style Setup ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    h_style = doc.styles[f"Heading {level}"]
    h_font = h_style.font
    h_font.name = "Calibri"
    h_font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    if level == 1:
        h_font.size = Pt(22)
        h_font.bold = True
    elif level == 2:
        h_font.size = Pt(16)
        h_font.bold = True
    elif level == 3:
        h_font.size = Pt(13)
        h_font.bold = True


def add_horizontal_line():
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="C7D2FE"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


def add_screenshot_placeholder(caption="[Insert Screenshot Here]"):
    """Add a bordered placeholder box for screenshots."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(5.5)

    # Set cell border and shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="dashed" w:sz="6" w:space="0" w:color="94A3B8"/>'
        '<w:left w:val="dashed" w:sz="6" w:space="0" w:color="94A3B8"/>'
        '<w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="94A3B8"/>'
        '<w:right w:val="dashed" w:sz="6" w:space="0" w:color="94A3B8"/>'
        "</w:tcBorders>"
    )
    tcPr.append(tcBorders)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC" w:val="clear"/>')
    tcPr.append(shading)

    # Add placeholder text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run(caption)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.italic = True

    # Add caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure: {caption.replace('[Insert Screenshot Here] — ', '')}")
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()


def add_styled_table(headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = ""
        p = header_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        # Purple background
        tc = header_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5" w:val="clear"/>')
        tcPr.append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            # Alternate row shading
            if r_idx % 2 == 1:
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2FF" w:val="clear"/>')
                tcPr.append(shading)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Spacer
for _ in range(4):
    doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Vault-AI")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
run.font.name = "Calibri"

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Privacy-Preserving Document Intelligence")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
run.font.name = "Calibri"

add_horizontal_line()

# Description
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.paragraph_format.space_before = Pt(12)
run = desc.add_run(
    "A privacy-first document intelligence system that automatically detects,\n"
    "masks, and encrypts personal information before LLM processing.\n"
    "Enabling secure question-answering without exposing sensitive data."
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.font.name = "Calibri"

# Spacer
for _ in range(6):
    doc.add_paragraph()

# Meta info
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Project Report")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("April 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# Page break
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)
add_horizontal_line()

toc_items = [
    ("1.", "Abstract"),
    ("2.", "Introduction"),
    ("3.", "Problem Statement"),
    ("4.", "System Architecture"),
    ("5.", "Technology Stack"),
    ("6.", "Implementation Details"),
    ("", "6.1  PII Detection Pipeline"),
    ("", "6.2  Tokenization Engine"),
    ("", "6.3  Encryption & Vault System"),
    ("", "6.4  Vector Store & Embeddings"),
    ("", "6.5  LLM Integration"),
    ("", "6.6  Differential Privacy"),
    ("7.", "Privacy & Security Model"),
    ("8.", "Application Walkthrough"),
    ("9.", "Comparison with Existing Solutions"),
    ("10.", "Limitations"),
    ("11.", "Future Work"),
    ("12.", "Conclusion"),
    ("13.", "References"),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f"{num}  {item}")
        run.font.size = Pt(12)
        run.font.bold = True
    else:
        run = p.add_run(f"      {item}")
        run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Abstract", level=1)
add_horizontal_line()

doc.add_paragraph(
    "Large Language Models (LLMs) have transformed document analysis and question-answering, "
    "but they pose significant privacy risks when processing sensitive documents containing "
    "Personally Identifiable Information (PII). Organizations and individuals are forced to "
    "choose between leveraging AI capabilities and protecting private data."
)
doc.add_paragraph(
    "Vault-AI addresses this challenge by introducing a privacy-preserving architecture that "
    "automatically detects and masks PII before any data reaches the LLM. The system employs "
    "Microsoft Presidio for Named Entity Recognition (NER), replaces sensitive values with "
    "deterministic token labels (e.g., PERSON_001, EMAIL_ADDRESS_001), and encrypts the "
    "token-to-value mappings using AES-256-GCM encryption with PBKDF2 key derivation. "
    "Document embeddings are further protected with differential privacy noise before storage "
    "in a vector database."
)
doc.add_paragraph(
    "The result is a system where users can upload sensitive documents, ask natural language "
    "questions, and receive accurate answers — while the LLM never sees any real personal "
    "information. Only the authenticated user, with their vault password, can reveal the "
    "original values in the responses."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Introduction", level=1)
add_horizontal_line()

doc.add_paragraph(
    "The rapid adoption of Large Language Models in enterprise and personal workflows has "
    "created an urgent need for privacy-preserving AI systems. When users upload contracts, "
    "medical records, financial statements, or HR documents to AI-powered tools, they "
    "inadvertently expose sensitive information — names, email addresses, phone numbers, "
    "social security numbers, credit card details, and more — to third-party API providers."
)
doc.add_paragraph(
    "Current LLM providers (OpenAI, Anthropic, Google) may log, store, or use submitted "
    "data for model training and debugging. Even with data processing agreements, the "
    "fundamental problem remains: raw PII leaves the user's control the moment it's sent "
    "to an API endpoint."
)
doc.add_paragraph(
    "Vault-AI takes a fundamentally different approach. Rather than trusting the LLM "
    "provider with sensitive data, the system creates a privacy firewall:"
)

items = [
    "PII is detected and replaced with deterministic tokens before the LLM ever sees the text",
    "The mapping between tokens and real values is AES-256 encrypted with a user-controlled password",
    "Document embeddings are protected with calibrated differential privacy noise",
    "Decryption happens locally — the LLM provider has zero access to original PII",
]
for item in items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph(
    "This report details the architecture, implementation, security model, and limitations "
    "of the Vault-AI system."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Problem Statement", level=1)
add_horizontal_line()

doc.add_paragraph(
    "The core problem Vault-AI addresses can be stated as follows:"
)

# Problem box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '"How can we leverage the power of Large Language Models for document intelligence '
    'while ensuring that no Personally Identifiable Information (PII) is ever exposed '
    'to the LLM provider or any unauthorized party?"'
)
run.font.italic = True
run.font.size = Pt(11.5)
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

doc.add_heading("Specific Challenges", level=2)

challenges = [
    ("PII Detection Accuracy", "Identifying diverse PII types (names, emails, phone numbers, "
     "financial data, dates, locations, IDs) across different document formats with high precision and recall."),
    ("Consistent Tokenization", "Replacing PII with deterministic labels that preserve document "
     "coherence — the same entity must always map to the same token for cross-document consistency."),
    ("Secure Key Management", "Encrypting token mappings with user-controlled passwords while "
     "ensuring usability — the system must never store passwords."),
    ("Embedding Privacy", "Preventing information leakage through vector embeddings, which "
     "can theoretically be reverse-engineered to recover original text fragments."),
    ("Query Accuracy", "Maintaining high-quality question-answering despite the text being "
     "tokenized — the LLM must understand context even with masked entities."),
]

for title, desc in challenges:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. System Architecture", level=1)
add_horizontal_line()

doc.add_paragraph(
    "Vault-AI follows a layered privacy-preserving architecture with clear separation "
    "between sensitive and sanitized data paths. The system consists of six core components "
    "that work together to ensure PII never reaches the LLM."
)

doc.add_heading("High-Level Architecture Diagram", level=2)

add_screenshot_placeholder("[Insert Screenshot Here] — System Architecture Diagram")

doc.add_heading("Component Overview", level=2)

add_styled_table(
    ["Component", "Technology", "Purpose"],
    [
        ["PII Detector", "Microsoft Presidio + spaCy NER", "Detect personal information in documents"],
        ["Tokenizer", "Custom deterministic engine", "Replace PII with consistent token labels"],
        ["Encryption Vault", "AES-256-GCM + PBKDF2", "Encrypt token-to-value mappings"],
        ["Vector Store", "ChromaDB + Sentence Transformers", "Store and query document embeddings"],
        ["Differential Privacy", "Laplacian noise mechanism", "Add calibrated noise to embeddings"],
        ["LLM Service", "Groq API (Llama-3)", "Generate answers from tokenized context"],
        ["Frontend", "Streamlit", "User interface for upload, query, and vault management"],
    ],
)

doc.add_heading("Data Flow", level=2)

doc.add_paragraph(
    "The following describes the complete data flow through the system:"
)

flow_steps = [
    ("Document Upload", "User uploads a PDF, TXT, or Markdown file. The DocumentProcessor extracts raw text."),
    ("PII Detection", "The PIIDetector uses Microsoft Presidio with spaCy's NER model to identify entities: "
     "PERSON, EMAIL_ADDRESS, PHONE_NUMBER, CREDIT_CARD, LOCATION, DATE_TIME, and more."),
    ("Tokenization", "Each detected PII entity is replaced with a deterministic label (e.g., "
     "'John Doe' → 'PERSON_001'). The same entity always maps to the same token across documents."),
    ("Vault Encryption", "The token-to-value mapping is encrypted using AES-256-GCM with a key "
     "derived from the user's password via PBKDF2 (480,000 iterations). Stored as vault.enc."),
    ("Embedding + DP Noise", "The tokenized text is chunked, embedded using Sentence Transformers, "
     "and Laplacian noise (calibrated by epsilon) is added for differential privacy."),
    ("Vector Storage", "Noisy embeddings and tokenized chunks are stored in ChromaDB locally."),
    ("Query Processing", "User's question is embedded and matched against stored chunks. "
     "Relevant tokenized context is sent to Llama-3 via Groq API."),
    ("Decryption", "User enters their vault password. The system decrypts the token mappings "
     "and replaces tokens in the answer with real values — all locally."),
]

for i, (title, desc) in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {i} — {title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Technology Stack", level=1)
add_horizontal_line()

add_styled_table(
    ["Layer", "Technology", "Version / Details"],
    [
        ["Frontend", "Streamlit", "Multi-page app with custom CSS"],
        ["PII Detection", "Microsoft Presidio", "Analyzer + spaCy NER (en_core_web_lg)"],
        ["Tokenization", "Custom Python Engine", "Deterministic label generation"],
        ["Encryption", "cryptography (Python)", "AES-256-GCM + PBKDF2-HMAC-SHA256"],
        ["Vector Database", "ChromaDB", "Persistent client with cosine similarity"],
        ["Embeddings", "Sentence Transformers", "all-MiniLM-L6-v2 (384-dim)"],
        ["Differential Privacy", "NumPy (Laplacian)", "Configurable epsilon parameter"],
        ["LLM", "Groq API", "Llama-3-8B-Instruct"],
        ["Experiment Tracking", "MLflow", "Local tracking for metrics & artifacts"],
        ["Language", "Python", "3.11+"],
    ],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. IMPLEMENTATION DETAILS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Implementation Details", level=1)
add_horizontal_line()

# 6.1 PII Detection
doc.add_heading("6.1  PII Detection Pipeline", level=2)

doc.add_paragraph(
    "The PII detection system uses Microsoft Presidio's AnalyzerEngine, which combines "
    "rule-based pattern matching with spaCy's Named Entity Recognition model. The system "
    "detects the following entity types:"
)

add_styled_table(
    ["Entity Type", "Examples", "Detection Method"],
    [
        ["PERSON", "John Doe, Dr. Smith", "spaCy NER + context rules"],
        ["EMAIL_ADDRESS", "john@example.com", "Regex pattern matching"],
        ["PHONE_NUMBER", "+1-555-0123, (555) 123-4567", "Regex + format validation"],
        ["CREDIT_CARD", "4532-1234-5678-9012", "Luhn algorithm + regex"],
        ["LOCATION", "New York, 123 Main St", "spaCy NER + gazetteer"],
        ["DATE_TIME", "January 15, 2024, 03/15/2024", "Pattern matching + NER"],
        ["US_SSN", "123-45-6789", "Regex with format validation"],
        ["IBAN_CODE", "GB82 WEST 1234 5698 7654 32", "Pattern + checksum"],
    ],
)

doc.add_paragraph(
    "Each detected entity includes a confidence score (0.0 to 1.0) indicating the "
    "model's certainty. Entities below a configurable threshold are excluded to minimize "
    "false positives."
)

# 6.2 Tokenization
doc.add_heading("6.2  Tokenization Engine", level=2)

doc.add_paragraph(
    "The tokenization engine replaces each PII entity with a deterministic, type-prefixed label. "
    "This design ensures:"
)

items = [
    "Consistency: The same entity always maps to the same token (e.g., every occurrence of "
    "'John Doe' becomes 'PERSON_001')",
    "Readability: Token labels are human-readable and preserve document structure",
    "Reversibility: The mapping can be reversed during decryption to restore original values",
    "Type Preservation: The token prefix (PERSON_, EMAIL_ADDRESS_) tells the LLM what kind of entity it is",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Example transformation:"
)

p = doc.add_paragraph()
run = p.add_run("Original: ")
run.bold = True
run = p.add_run('"John Doe (john@gmail.com) submitted the Q4 financial report on March 15, 2024."')
run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

p = doc.add_paragraph()
run = p.add_run("Tokenized: ")
run.bold = True
run = p.add_run('"PERSON_001 (EMAIL_ADDRESS_001) submitted the Q4 financial report on DATE_TIME_001."')
run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

# 6.3 Encryption
doc.add_heading("6.3  Encryption & Vault System", level=2)

doc.add_paragraph(
    "The vault stores the token-to-value mapping encrypted with the user's password. "
    "The encryption uses industry-standard algorithms:"
)

add_styled_table(
    ["Parameter", "Value", "Purpose"],
    [
        ["Cipher", "AES-256-GCM", "Authenticated encryption with associated data"],
        ["Key Derivation", "PBKDF2-HMAC-SHA256", "Derive encryption key from password"],
        ["Iterations", "480,000", "Resist brute-force password cracking"],
        ["Salt", "16 bytes (random per encryption)", "Prevent rainbow table attacks"],
        ["Nonce", "12 bytes (random per encryption)", "Ensure unique ciphertext each time"],
        ["Key Length", "32 bytes (256 bits)", "Maximum AES key strength"],
    ],
)

doc.add_paragraph(
    "The encrypted vault file format is: [16 bytes salt] [12 bytes nonce] [ciphertext + GCM tag]. "
    "Each encryption operation generates a unique salt and nonce, ensuring that encrypting "
    "the same data twice produces completely different ciphertext."
)

# 6.4 Vector Store
doc.add_heading("6.4  Vector Store & Embeddings", level=2)

doc.add_paragraph(
    "Document chunks are embedded using the all-MiniLM-L6-v2 sentence transformer model "
    "(384-dimensional vectors) and stored in ChromaDB with cosine similarity indexing. "
    "Only the tokenized (masked) text is embedded — the original text with PII is never "
    "stored in the vector database."
)

# 6.5 LLM Integration
doc.add_heading("6.5  LLM Integration", level=2)

doc.add_paragraph(
    "Vault-AI uses Groq's API to access Meta's Llama-3-8B-Instruct model. The LLM receives "
    "only tokenized context and generates answers using the token labels. The system prompt "
    "instructs the model to treat tokens as entity references and provide coherent answers."
)

p = doc.add_paragraph()
run = p.add_run("Key point: ")
run.bold = True
run = p.add_run(
    "The LLM never sees real PII. It processes text like 'PERSON_001 submitted the report' "
    "and generates answers like 'According to the document, PERSON_001 was responsible.' "
    "The actual identity is only revealed after local decryption."
)

# 6.6 Differential Privacy
doc.add_heading("6.6  Differential Privacy", level=2)

doc.add_paragraph(
    "To protect against embedding inversion attacks (where an adversary tries to reconstruct "
    "original text from embedding vectors), Vault-AI adds calibrated Laplacian noise to all "
    "embeddings before storage."
)

doc.add_paragraph(
    "The noise is controlled by the epsilon (ε) parameter — a standard differential privacy "
    "parameter where lower values provide stronger privacy guarantees but may slightly reduce "
    "retrieval accuracy. The default epsilon is calibrated to balance privacy and utility."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. PRIVACY & SECURITY MODEL
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Privacy & Security Model", level=1)
add_horizontal_line()

doc.add_heading("Data Classification", level=2)

add_styled_table(
    ["Data", "Storage Location", "Format", "Risk if Exposed"],
    [
        ["Original text with PII", "Nowhere (discarded)", "N/A — never persisted", "N/A"],
        ["Token mappings", "vault_data/vault.enc", "AES-256-GCM encrypted", "Low — requires password"],
        ["Tokenized document text", "chroma_data/ (ChromaDB)", "Masked text (PERSON_001, etc.)", "Minimal — no real PII"],
        ["Document embeddings", "chroma_data/ (ChromaDB)", "Vectors + DP noise", "Minimal — noised"],
        ["User password", "Browser session (RAM only)", "Never persisted to disk", "N/A"],
    ],
)

doc.add_heading("Threat Analysis", level=2)

add_styled_table(
    ["Attack Scenario", "Traditional LLM", "Vault-AI"],
    [
        ["LLM provider data breach", "All PII exposed in plain text", "Only tokenized text — no real PII"],
        ["API traffic interception", "PII visible in requests", "Only token labels in requests"],
        ["Server filesystem breach", "PII in database in plain text", "Encrypted vault + masked text only"],
        ["Server breach + password theft", "Full exposure", "Full exposure (same risk)"],
        ["Embedding inversion attack", "Original text recoverable", "DP noise prevents reconstruction"],
        ["Insider threat at LLM provider", "Can read all user data", "Can only see PERSON_001, etc."],
    ],
)

doc.add_paragraph(
    "The security model provides defense-in-depth: even if one layer is compromised, "
    "the remaining layers continue to protect user data. The most common real-world threats "
    "(LLM provider breaches, API logging, traffic interception) are fully mitigated."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. APPLICATION WALKTHROUGH
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Application Walkthrough", level=1)
add_horizontal_line()

doc.add_heading("8.1  Home Page — Password Setup", level=2)
doc.add_paragraph(
    "The home page provides an overview of Vault-AI and allows users to create their "
    "vault password. This password is used for all encryption and decryption operations."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Home Page with Password Setup")

doc.add_heading("8.2  Upload Page — PII Detection", level=2)
doc.add_paragraph(
    "Users upload documents (PDF, TXT, Markdown) and the system automatically detects "
    "PII entities. The preview shows detected entities as token labels — real values are "
    "never displayed."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Document Upload & PII Detection")

doc.add_heading("8.3  Upload Page — Encryption & Storage", level=2)
doc.add_paragraph(
    "After PII detection, users enter their vault password to encrypt the token mappings "
    "and store the document securely. The progress bar shows each step of the pipeline."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Encryption & Storage Progress")

doc.add_heading("8.4  Query Page — Privacy Mode", level=2)
doc.add_paragraph(
    "Users ask natural language questions about their documents. Answers are displayed "
    "in Privacy Mode with token labels (PERSON_001, etc.) instead of real values."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Query Results in Privacy Mode")

doc.add_heading("8.5  Query Page — Decrypted View", level=2)
doc.add_paragraph(
    "By entering their vault password, users can decrypt answers to reveal real values. "
    "Decrypted entities are highlighted in green for easy identification."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Decrypted Query Results")

doc.add_heading("8.6  Vault Page — Statistics", level=2)
doc.add_paragraph(
    "The vault page shows statistics about stored documents, PII entity counts, and "
    "entity type distribution with interactive charts."
)
add_screenshot_placeholder("[Insert Screenshot Here] — Vault Statistics & Charts")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. COMPARISON WITH EXISTING SOLUTIONS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Comparison with Existing Solutions", level=1)
add_horizontal_line()

add_styled_table(
    ["Feature", "ChatGPT / Claude", "Microsoft Purview", "Vault-AI"],
    [
        ["PII leaves user's control", "Yes — sent to API", "Partial — cloud-based", "No — masked before API call"],
        ["Encryption of PII mappings", "None", "Platform-managed keys", "User-controlled AES-256"],
        ["Differential privacy on embeddings", "No", "No", "Yes (Laplacian noise)"],
        ["Password-based decryption", "N/A", "Admin-controlled", "User-controlled per session"],
        ["LLM sees real PII", "Yes", "Depends on config", "Never"],
        ["Open source", "No", "No", "Yes"],
        ["Self-hostable", "No", "No", "Yes"],
        ["Works offline (except LLM call)", "No", "No", "Yes"],
    ],
)

doc.add_paragraph(
    "Vault-AI's key differentiator is that it treats the LLM as an untrusted component. "
    "While other solutions may offer data governance policies or contractual protections, "
    "Vault-AI provides cryptographic guarantees that PII never reaches the LLM."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Limitations", level=1)
add_horizontal_line()

limitations = [
    ("PII Detection is Not Perfect",
     "Microsoft Presidio and spaCy NER may miss some PII entities (false negatives) or "
     "incorrectly flag non-PII text (false positives). Domain-specific entities (e.g., "
     "internal project codes, medical codes) may require custom recognizers."),

    ("Server-Side Deployment Risks",
     "When deployed on a server, the encrypted vault file and ChromaDB data reside on "
     "that server. An attacker who gains both server access AND the user's password could "
     "decrypt the vault. The system does not implement client-side encryption."),

    ("Single Password Model",
     "The system uses a single password for all encryption. There is no multi-user "
     "authentication, role-based access control, or password recovery mechanism. "
     "If the user forgets their password, encrypted data is irrecoverable."),

    ("Context Loss from Tokenization",
     "Replacing names and entities with tokens can reduce the LLM's ability to leverage "
     "real-world knowledge about those entities. For example, the LLM cannot use its "
     "knowledge about a specific company if the company name is tokenized."),

    ("No Support for Images or Tables",
     "The current system only processes text content. PII embedded in images, scanned "
     "documents (without OCR), or complex table structures may not be detected."),

    ("Embedding Privacy is Approximate",
     "Differential privacy noise on embeddings provides mathematical privacy guarantees, "
     "but higher noise levels reduce retrieval accuracy. The epsilon parameter requires "
     "careful tuning for each use case."),

    ("Single LLM Provider Dependency",
     "The system currently relies on Groq's API for LLM inference. While the tokenized "
     "text is sent (not real PII), the system still depends on an external API for "
     "question-answering functionality."),
]

for title, desc in limitations:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. FUTURE WORK
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("11. Future Work", level=1)
add_horizontal_line()

future_items = [
    ("Client-Side Encryption (Zero-Knowledge Architecture)",
     "Implement encryption/decryption entirely in the browser using the WebCrypto API. "
     "The server would store encrypted blobs but never have access to the encryption key "
     "or decrypted data, achieving true zero-knowledge privacy."),

    ("Multi-User Authentication & RBAC",
     "Add proper user authentication (OAuth2, SSO) with role-based access control. "
     "Each user would have their own encrypted vault with separate permissions for "
     "viewing, querying, and managing documents."),

    ("Custom PII Recognizers",
     "Allow users to define custom entity types and recognition patterns for "
     "domain-specific PII (e.g., patient IDs in healthcare, employee codes in HR, "
     "case numbers in legal documents)."),

    ("Local LLM Support",
     "Integrate local LLM inference (e.g., Ollama, llama.cpp) to eliminate the need "
     "for any external API calls, achieving fully offline operation with complete "
     "data sovereignty."),

    ("OCR & Multimodal PII Detection",
     "Add OCR (Optical Character Recognition) support for scanned documents and image-based "
     "PDFs. Extend PII detection to images, tables, and handwritten text."),

    ("Audit Logging & Compliance",
     "Implement comprehensive audit trails for all encryption, decryption, and query "
     "operations to support compliance requirements (GDPR, HIPAA, CCPA)."),

    ("Hardware Security Module (HSM) Integration",
     "Support cloud-based key management services (AWS KMS, Azure Key Vault, Google Cloud KMS) "
     "for enterprise-grade key management and hardware-backed encryption."),

    ("Federated Learning for PII Detection",
     "Improve PII detection accuracy over time using federated learning — the model learns "
     "from corrections across users without any user's data leaving their environment."),
]

for title, desc in future_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("12. Conclusion", level=1)
add_horizontal_line()

doc.add_paragraph(
    "Vault-AI demonstrates that privacy and AI utility are not mutually exclusive. "
    "By introducing a privacy firewall between sensitive documents and the LLM, the "
    "system enables secure document intelligence without compromising on user experience "
    "or answer quality."
)
doc.add_paragraph(
    "The architecture — combining NER-based PII detection, deterministic tokenization, "
    "AES-256-GCM encryption, differential privacy, and retrieval-augmented generation — "
    "provides multiple layers of defense. The most common and impactful threats "
    "(LLM provider data breaches, API traffic interception, insider threats) are "
    "cryptographically eliminated."
)
doc.add_paragraph(
    "While limitations exist (particularly around server-side deployment and PII detection "
    "accuracy), the system represents a significant step toward privacy-preserving AI. "
    "The open-source, self-hostable nature of the project enables organizations and "
    "individuals to take control of their data privacy while still benefiting from "
    "state-of-the-art language model capabilities."
)
doc.add_paragraph(
    "As AI systems become increasingly integrated into workflows handling sensitive data, "
    "architectures like Vault-AI — where the AI works on sanitized data and users control "
    "the decryption keys — will become not just desirable, but essential."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 13. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("13. References", level=1)
add_horizontal_line()

references = [
    "Microsoft Presidio — Data Protection and De-identification SDK. https://github.com/microsoft/presidio",
    "spaCy — Industrial-Strength Natural Language Processing. https://spacy.io/",
    "ChromaDB — Open-Source Embedding Database. https://www.trychroma.com/",
    "Sentence Transformers — Multilingual Sentence Embeddings. https://www.sbert.net/",
    "Groq — Fast AI Inference Platform. https://groq.com/",
    "Meta Llama 3 — Open Foundation Language Models. https://ai.meta.com/llama/",
    "Dwork, C. (2006). 'Differential Privacy.' ICALP 2006.",
    "NIST SP 800-132 — Recommendation for Password-Based Key Derivation. https://csrc.nist.gov/",
    "AES-GCM — Galois/Counter Mode of Operation. NIST SP 800-38D.",
    "MLflow — Open Source Platform for ML Lifecycle. https://mlflow.org/",
    "Streamlit — The Fastest Way to Build Data Apps. https://streamlit.io/",
    "GDPR — General Data Protection Regulation. https://gdpr.eu/",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}]  ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "Vault-AI_Project_Report.docx")
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
